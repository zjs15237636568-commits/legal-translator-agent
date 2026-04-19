"""
中英对照 Word 导出（双栏表格 + 风险批注）。

MVP：使用两列表格 (English | 中文)，章节前插入标题。风险以尾注形式附在文档末尾。
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm

from app.core.db import get_db


LEVEL_LABEL = {"red": "🔴 法律风险", "yellow": "🟡 业务提示"}


async def export_bilingual_docx(project_id: str) -> bytes:
    async with get_db() as db:
        p = await (await db.execute(
            "SELECT * FROM projects WHERE id = ?", (project_id,)
        )).fetchone()
        if not p:
            raise ValueError("Project not found")
        segs = await (await db.execute(
            "SELECT s.id, s.seq, s.clause_no, s.heading, s.original_md, "
            "       t.translated_md "
            "FROM segments s LEFT JOIN translations t ON t.segment_id = s.id "
            "WHERE s.project_id = ? ORDER BY s.seq ASC",
            (project_id,),
        )).fetchall()
        risks = await (await db.execute(
            "SELECT * FROM risks WHERE project_id = ? "
            "ORDER BY CASE level WHEN 'red' THEN 0 ELSE 1 END, created_at ASC",
            (project_id,),
        )).fetchall()

    doc = Document()
    # 文档标题
    title = doc.add_heading(p["name"] or "Bilingual Contract", level=0)
    doc.add_paragraph(f"Source: {p['source_name']}").italic = True

    # 表格：English | 中文
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "English (Original)"
    hdr[1].text = "中文（译文）"
    for c in hdr:
        for p_ in c.paragraphs:
            for r_ in p_.runs:
                r_.bold = True

    for s in segs:
        row = table.add_row().cells
        row[0].text = s["original_md"] or ""
        row[1].text = s["translated_md"] or ""
        for c in row:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    # 风险附录
    if risks:
        doc.add_page_break()
        doc.add_heading("风险审查 Risk Review", level=1)
        for r in risks:
            label = LEVEL_LABEL.get(r["level"], r["level"])
            p_h = doc.add_paragraph()
            run = p_h.add_run(f"{label} · {r['category']} · {r['title']}")
            run.bold = True
            run.font.size = Pt(11)
            doc.add_paragraph(f"详情：{r['detail']}")
            doc.add_paragraph(f"建议：{r['suggestion']}")
            doc.add_paragraph("")  # 空行

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
